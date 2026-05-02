#!/usr/bin/env python3
"""
generate_docx.py — Generate professional .docx resumes and cover letters.

Usage:
  python scripts/generate_docx.py --input content.json --output output.docx [--type resume|cover_letter]

The input JSON must match the schema documented below. This script is called by
the resume-writer and cover-letter agents. All document generation MUST use this
script — agents should write structured JSON and call it, never generate inline code.

Requires: pip install python-docx
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('Error: "python-docx" package not found. Install it with: pip install python-docx',
          file=sys.stderr)
    sys.exit(1)


# ── Constants ────────────────────────────────────────────────────────────────

GRAY = RGBColor(0x66, 0x66, 0x66)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _hex_color(hex_str: str) -> RGBColor:
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _clear_default_paragraph(doc: Document) -> None:
    """Remove the empty paragraph python-docx inserts at document creation."""
    if doc.paragraphs:
        elem = doc.paragraphs[0]._element
        elem.getparent().remove(elem)


def _new_para(doc: Document, align=None, space_before: float = 0,
              space_after: float = 0, left_indent=None, first_line_indent=None,
              line_spacing=None):
    """
    Append a blank paragraph with explicit layout. Caller adds runs via _run().
    All spacing values are in points (converted from twips where applicable).
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        para.alignment = align
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return para


def _run(para, text: str, font_name: str = 'Calibri', size_pt: float = 11,
         bold: bool = False, italic: bool = False, color: RGBColor = None) -> None:
    """Append a formatted text run to an existing paragraph."""
    if not text:
        return
    r = para.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    if bold:
        r.font.bold = True
    if italic:
        r.font.italic = True
    if color is not None:
        r.font.color.rgb = color


def _bottom_border(para, hex_color: str) -> None:
    """
    Add a single hairline bottom border to a paragraph used as a section divider.
    hex_color: 6-char hex string without '#', e.g. '2B579A'
    Border thickness w:sz=6 → 0.75pt; w:space=4 → 4pt gap between text and line.
    """
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc, text: str, font: str, heading_size_pt: float,
                     accent_rgb: RGBColor, accent_hex: str):
    """
    Section header: ALL CAPS, bold, accent color, 12pt before / 4pt after,
    with a single hairline bottom border as a divider.
    Spacing: 240 twips before (12pt), 80 twips after (4pt).
    """
    para = _new_para(doc, space_before=12, space_after=4)
    _run(para, text.upper(), font_name=font, size_pt=heading_size_pt,
         bold=True, color=accent_rgb)
    _bottom_border(para, accent_hex)
    return para


# ── Page setup ───────────────────────────────────────────────────────────────

def _configure_page(doc: Document, style: dict) -> None:
    """Set Letter page size and configurable margins from style dict."""
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(style.get('margin_top', 0.75))
    sec.bottom_margin = Inches(style.get('margin_bottom', 0.75))
    sec.left_margin = Inches(style.get('margin_left', 0.85))
    sec.right_margin = Inches(style.get('margin_right', 0.85))


def _configure_cover_letter_page(doc: Document) -> None:
    """Cover letters use 1-inch margins on all sides."""
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)


# ── Resume builder ───────────────────────────────────────────────────────────

def _build_resume(doc: Document, content: dict) -> None:
    """
    Populate doc with resume content.

    Expected content keys (all optional except where noted):
      personal: { name*, title, email, phone, linkedin, location }
      summary: str
      skills: { category: [skill, ...] }
      experience: [{ company, title, dates, location, bullets: [...] }]
      education: [{ institution, degree, field, date, honors }]
      certifications: [{ name, issuer, date }]
      publications: [{ title, venue, date, summary }]
      speaking: [{ title, event, date, description }]
      projects: [{ name, technologies, description, highlights: [...] }]
      awards: [{ name, issuer, description }]
      patents: [{ title, number, date, description }]
      style: { font, accent_color, body_size_pt, heading_size_pt, name_size_pt,
               margin_top, margin_bottom, margin_left, margin_right }
    """
    s = content.get('style', {})
    font = s.get('font', 'Calibri')
    body = s.get('body_size_pt', 11)
    heading = s.get('heading_size_pt', 13)
    name_sz = s.get('name_size_pt', 18)
    accent_hex = s.get('accent_color', '2B579A')
    accent_rgb = _hex_color(accent_hex)
    p = content.get('personal', {})

    # ── Header ───────────────────────────────────────────────────────────────

    # Name — centered, bold, 18pt (default), space after 2pt (40 twips)
    para = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _run(para, p.get('name', 'Candidate Name'), font_name=font, size_pt=name_sz, bold=True)

    # Title — centered, accent color, body+1pt, space after 2pt
    if p.get('title'):
        para = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
        _run(para, p['title'], font_name=font, size_pt=body + 1, color=accent_rgb)

    # Contact line — centered, body-1pt, pipe-separated, space after 10pt (200 twips)
    contact_parts = [p.get('email'), p.get('phone'), p.get('location'), p.get('linkedin')]
    contact = '  |  '.join(x for x in contact_parts if x)
    if contact:
        para = _new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
        _run(para, contact, font_name=font, size_pt=body - 1)

    # ── Summary ──────────────────────────────────────────────────────────────

    if content.get('summary'):
        _section_heading(doc, 'Professional Summary', font, heading, accent_rgb, accent_hex)
        # 6pt after (120 twips)
        para = _new_para(doc, space_after=6)
        _run(para, content['summary'], font_name=font, size_pt=body)

    # ── Skills ───────────────────────────────────────────────────────────────

    skills = content.get('skills') or {}
    if skills:
        _section_heading(doc, 'Technical Skills', font, heading, accent_rgb, accent_hex)
        for category, items in skills.items():
            para = _new_para(doc, space_after=3)  # 60 twips
            _run(para, f'{category}: ', font_name=font, size_pt=body, bold=True)
            _run(para, ', '.join(items), font_name=font, size_pt=body)

    # ── Experience ───────────────────────────────────────────────────────────

    experience = content.get('experience') or []
    if experience:
        _section_heading(doc, 'Professional Experience', font, heading, accent_rgb, accent_hex)
        for role in experience:
            # Title — Company: bold title, regular company, 8pt before / 1pt after (160/20 twips)
            para = _new_para(doc, space_before=8, space_after=1)
            _run(para, role.get('title', ''), font_name=font, size_pt=body, bold=True)
            if role.get('company'):
                _run(para, f"  —  {role['company']}", font_name=font, size_pt=body)

            # Dates | Location: italic, gray, body-1pt, 3pt after (60 twips)
            meta_parts = [role.get('dates'), role.get('location')]
            meta = '  |  '.join(x for x in meta_parts if x)
            if meta:
                para = _new_para(doc, space_after=3)
                _run(para, meta, font_name=font, size_pt=body - 1, italic=True, color=GRAY)

            # Bullets: hanging indent left=18pt/first=-9pt (360/180 twips), 2pt after (40 twips)
            for bullet in role.get('bullets', []):
                para = _new_para(doc, space_after=2, left_indent=18, first_line_indent=-9)
                _run(para, f'•  {bullet}', font_name=font, size_pt=body)

    # ── Certifications ───────────────────────────────────────────────────────
    # Order matches resume_style.yaml section_order: certifications before education

    certifications = content.get('certifications') or []
    if certifications:
        _section_heading(doc, 'Certifications', font, heading, accent_rgb, accent_hex)
        for cert in certifications:
            para = _new_para(doc, space_after=2)
            _run(para, cert.get('name', ''), font_name=font, size_pt=body, bold=True)
            if cert.get('issuer'):
                _run(para, f"  —  {cert['issuer']}", font_name=font, size_pt=body)
            if cert.get('date'):
                _run(para, f"  ({cert['date']})", font_name=font, size_pt=body - 1, color=GRAY)

    # ── Education ────────────────────────────────────────────────────────────

    education = content.get('education') or []
    if education:
        _section_heading(doc, 'Education', font, heading, accent_rgb, accent_hex)
        for edu in education:
            para = _new_para(doc, space_after=3)
            _run(para, f"{edu.get('degree', '')} in {edu.get('field', '')}",
                 font_name=font, size_pt=body, bold=True)
            if edu.get('institution'):
                _run(para, f"  —  {edu['institution']}", font_name=font, size_pt=body)
            if edu.get('date'):
                _run(para, f"  ({edu['date']})", font_name=font, size_pt=body - 1, color=GRAY)
            if edu.get('honors'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, edu['honors'], font_name=font, size_pt=body - 1, italic=True)

    # ── Publications ─────────────────────────────────────────────────────────

    publications = content.get('publications') or []
    if publications:
        _section_heading(doc, 'Publications', font, heading, accent_rgb, accent_hex)
        for pub in publications:
            para = _new_para(doc, space_after=2)
            _run(para, pub.get('title', ''), font_name=font, size_pt=body, bold=True)
            if pub.get('venue'):
                _run(para, f"  —  {pub['venue']}", font_name=font, size_pt=body)
            if pub.get('date'):
                _run(para, f"  ({pub['date']})", font_name=font, size_pt=body - 1, color=GRAY)
            if pub.get('summary'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, pub['summary'], font_name=font, size_pt=body - 1, italic=True)

    # ── Speaking ─────────────────────────────────────────────────────────────

    speaking = content.get('speaking') or []
    if speaking:
        _section_heading(doc, 'Speaking', font, heading, accent_rgb, accent_hex)
        for talk in speaking:
            para = _new_para(doc, space_after=2)
            _run(para, talk.get('title', ''), font_name=font, size_pt=body, bold=True)
            if talk.get('event'):
                _run(para, f"  —  {talk['event']}", font_name=font, size_pt=body)
            if talk.get('date'):
                _run(para, f"  ({talk['date']})", font_name=font, size_pt=body - 1, color=GRAY)
            if talk.get('description'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, talk['description'], font_name=font, size_pt=body - 1, italic=True)

    # ── Projects ─────────────────────────────────────────────────────────────

    projects = content.get('projects') or []
    if projects:
        _section_heading(doc, 'Projects', font, heading, accent_rgb, accent_hex)
        for proj in projects:
            # Name (technologies): 5pt before / 1pt after (100/20 twips)
            para = _new_para(doc, space_before=5, space_after=1)
            _run(para, proj.get('name', ''), font_name=font, size_pt=body, bold=True)
            if proj.get('technologies'):
                _run(para, f"  ({proj['technologies']})", font_name=font,
                     size_pt=body - 1, color=GRAY)
            if proj.get('description'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, proj['description'], font_name=font, size_pt=body)
            for highlight in proj.get('highlights', []):
                para = _new_para(doc, space_after=2, left_indent=18, first_line_indent=-9)
                _run(para, f'•  {highlight}', font_name=font, size_pt=body)

    # ── Awards ───────────────────────────────────────────────────────────────

    awards = content.get('awards') or []
    if awards:
        _section_heading(doc, 'Awards', font, heading, accent_rgb, accent_hex)
        for award in awards:
            para = _new_para(doc, space_after=2)
            _run(para, award.get('name', ''), font_name=font, size_pt=body, bold=True)
            if award.get('issuer'):
                _run(para, f"  —  {award['issuer']}", font_name=font, size_pt=body)
            if award.get('description'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, award['description'], font_name=font, size_pt=body - 1, italic=True)

    # ── Patents ──────────────────────────────────────────────────────────────

    patents = content.get('patents') or []
    if patents:
        _section_heading(doc, 'Patents', font, heading, accent_rgb, accent_hex)
        for patent in patents:
            para = _new_para(doc, space_after=2)
            _run(para, patent.get('title', ''), font_name=font, size_pt=body, bold=True)
            if patent.get('number'):
                _run(para, f"  —  {patent['number']}", font_name=font, size_pt=body)
            if patent.get('date'):
                _run(para, f"  ({patent['date']})", font_name=font, size_pt=body - 1, color=GRAY)
            if patent.get('description'):
                para = _new_para(doc, space_after=2, left_indent=18)
                _run(para, patent['description'], font_name=font, size_pt=body - 1, italic=True)


# ── Cover letter builder ─────────────────────────────────────────────────────

def _build_cover_letter(doc: Document, content: dict) -> None:
    """
    Populate doc with cover letter content.

    Expected content keys:
      personal: { name, email, phone, location }
      date: str  (e.g. 'February 13, 2026')
      recipient: { name, title, company, address }
      salutation: str  (default: 'Dear Hiring Team,')
      paragraphs: [str, ...]
      closing: str  (default: 'Sincerely,')
      style: { font, body_size_pt }
    """
    s = content.get('style', {})
    font = s.get('font', 'Calibri')
    body = s.get('body_size_pt', 11)
    p = content.get('personal', {})
    r = content.get('recipient', {})

    # Sender info — each line 1pt after (20 twips)
    for line in [p.get('name'), p.get('email'), p.get('phone'), p.get('location')]:
        if line:
            para = _new_para(doc, space_after=1)
            _run(para, line, font_name=font, size_pt=body)

    # Spacer — 10pt (200 twips)
    _new_para(doc, space_after=10)

    # Date — 10pt after
    if content.get('date'):
        para = _new_para(doc, space_after=10)
        _run(para, content['date'], font_name=font, size_pt=body)

    # Recipient — 1pt per line, then 10pt spacer
    recipient_lines = [r.get('name'), r.get('title'), r.get('company'), r.get('address')]
    recipient_lines = [x for x in recipient_lines if x]
    for line in recipient_lines:
        para = _new_para(doc, space_after=1)
        _run(para, line, font_name=font, size_pt=body)
    if recipient_lines:
        _new_para(doc, space_after=10)

    # Salutation — 10pt after
    para = _new_para(doc, space_after=10)
    _run(para, content.get('salutation', 'Dear Hiring Team,'), font_name=font, size_pt=body)

    # Body paragraphs — 10pt after (200 twips), line spacing 13.8pt exact (276 twips)
    for text in content.get('paragraphs', []):
        para = _new_para(doc, space_after=10, line_spacing=13.8)
        _run(para, text, font_name=font, size_pt=body)

    # Closing — 10pt before / 30pt after (200/600 twips)
    para = _new_para(doc, space_before=10, space_after=30)
    _run(para, content.get('closing', 'Sincerely,'), font_name=font, size_pt=body)

    # Signature
    para = _new_para(doc)
    _run(para, p.get('name', ''), font_name=font, size_pt=body)


# ── Entry point ──────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='Generate a professional .docx resume or cover letter from structured JSON.'
    )
    parser.add_argument('--input', required=True, help='Path to JSON content file')
    parser.add_argument('--output', required=True, help='Output path for .docx file')
    parser.add_argument('--type', default='resume', choices=['resume', 'cover_letter'],
                        help='Document type (default: resume)')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        content = json.load(f)

    doc = Document()
    _clear_default_paragraph(doc)

    if args.type == 'cover_letter':
        _configure_cover_letter_page(doc)
        _build_cover_letter(doc, content)
    else:
        _configure_page(doc, content.get('style', {}))
        _build_resume(doc, content)

    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    size_kb = out.stat().st_size / 1024
    print(f'Generated: {out} ({size_kb:.1f} KB)')


if __name__ == '__main__':
    main()
